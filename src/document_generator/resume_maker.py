import os
import subprocess
from typing import Dict, Tuple
from docx import Document
from datetime import datetime
import pandas as pd
from src.config import RESUME_TEMPLATES_DIR, OUTPUT_RESUMES_DIR

class ResumeManager:
    def __init__(self, df: pd.DataFrame):
        self.df = df
        self.process_all_resumes()

    def process_all_resumes(self) -> None:
        for _, row in self.df.iterrows():
            self.create_resume_and_cover(row)

    def find_resume_cover_template(self, job_category: str) -> Tuple[str, str]:
        resume_path = os.path.join(RESUME_TEMPLATES_DIR, f"resume_data role.docx")
        cover_path = os.path.join(RESUME_TEMPLATES_DIR, f"cover_data role.docx")
        try:
            if job_category:
                resume_path = os.path.join(RESUME_TEMPLATES_DIR, f"resume_{job_category}.docx")
                cover_path = os.path.join(RESUME_TEMPLATES_DIR, f"cover_{job_category}.docx")
            return resume_path, cover_path
        except Exception as e:
            print(f"Error finding template: {str(e)}")
            return resume_path, cover_path

    @staticmethod
    def table_edit_replace(doc: Document, target: str, value: str) -> None:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if target in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.text = run.text.replace(target, value)

    @staticmethod
    def paragraph_edit_replace(doc: Document, target: str, value: str) -> None:
        for paragraph in doc.paragraphs:
            if target in paragraph.text:
                for run in paragraph.runs:
                    if target in run.text:
                        run.text = run.text.replace(target, value)

    def create_resume_and_cover(self, row: pd.Series) -> None:
        resume_path, cover_path = self.find_resume_cover_template(row.get("job_category", ""))
        
        resume_doc = Document(resume_path)
        cover_doc = Document(cover_path)
        
        company = row.get("company_name", "")
        job_role = row.get("job_position_title", "Data Analyst")
        top_skills = row.get("top_skills", "Python, SQL, Power BI, Excel, Machine Learning")
        location = row.get("location", "Montreal, Canada")
        
        today_date = datetime.now().strftime("%d-%b-%Y")
        output_path = os.path.join(OUTPUT_RESUMES_DIR, f"{company}_{job_role}_{today_date}")
        os.makedirs(output_path, exist_ok=True)
        
        cover_input_dict = {
            "[job role]": job_role,
            "[company name]": company,
            "[company location]": row.get("location", ""),
            "[date]": today_date,
            "[why company]": row.get("why_this_company", ""),
            "[why me]": row.get("why_me", ""),
            "[location]": row.get("location", "")
        }
        
        self.create_resume(resume_doc, job_role, top_skills, location, output_path)
        self.create_cover(cover_doc, output_path, cover_input_dict)

    def create_cover(self, cover_doc: Document, output_path: str, cover_input_dict: Dict[str, str]) -> None:
        for key, value in cover_input_dict.items():
            self.table_edit_replace(cover_doc, key, value)
            self.paragraph_edit_replace(cover_doc, key, value)
        
        docx_path = os.path.join(output_path, "Krishnakumar Cover Letter.docx")
        cover_doc.save(docx_path)
        self.save_to_pdf(output_path, docx_path)

    def create_resume(self, resume_doc: Document, job_role: str, top_skills: str, location: str, output_path: str) -> None:
        self.table_edit_replace(resume_doc, "[job role]", job_role)
        self.table_edit_replace(resume_doc, "[location]", location)
        self.paragraph_edit_replace(resume_doc, "[top skills]", top_skills)
        docx_path = os.path.join(output_path, "Krishnakumar Resume.docx")
        resume_doc.save(docx_path)
        self.save_to_pdf(output_path, docx_path)

    @staticmethod
    def save_to_pdf(output_path: str, docx_path: str) -> None:
        libreoffice_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        try:
            subprocess.run([
                libreoffice_path,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                output_path,
                docx_path
            ], check=True)
        except subprocess.CalledProcessError as e:
            print(f"Error converting to PDF: {e}")

if __name__ == "__main__":
    # Add any test or example usage here
    pass
