import re
from docx import Document
from datetime import datetime
import pandas as pd
import os
import sys
import subprocess
import re

# current_dir = os.path.dirname(os.path.abspath(__file__))
# sys.path.insert(0, current_dir)


class ResumeManager():
    def __init__(self, df) -> None:
        self.df = df
        for index, row in self.df.iterrows():
            self.edit_resume_cover(row)
    
    def find_resume_cover_template(self, row):
        resume_path = f"ResumeManager/ResumeTemplates/resume_data role.docx"
        cover_path = f"ResumeManager/ResumeTemplates/cover_data role.docx"
        try:
            job_category = row["job_category"]
            if job_category:
                resume_path = f"ResumeManager/ResumeTemplates/resume_{job_category}.docx"
                cover_path = f"ResumeManager/ResumeTemplates/cover_{job_category}.docx"
            return resume_path, cover_path
        except Exception as e:
            print(f"An error occurred: {str(e)}")
            return resume_path, cover_path
    
    def table_edit_replace(self, doc, target, value):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if target in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.text = run.text.replace(target, value)
    
    def paragraph_edit_replace(self, doc, target, value):
        for paragraph in doc.paragraphs:
            if target in paragraph.text:
                inline = paragraph.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if target in inline[i].text:
                        text = inline[i].text.replace(target, value)
                        inline[i].text = text

    
    def edit_resume_cover(self, row):
        # find resume and cover letter template path based on job category
        resume_path, cover_path = self.find_resume_cover_template(row)
        
        resume_doc = Document(resume_path)
        cover_doc = Document(cover_path)
        top_skills = row["top_skills"]
        company = row["company_name"]
        job_role = row["job_position_title"]
        company_location = row["location"]
        today_date = datetime.now().strftime("%d-%b-%Y")
        why_company = row["why_this_company"]
        why_me = row["why_me"]
        
        if not job_role:
            job_role = "Data Analyst"
        if not top_skills:
            top_skills = "Python, SQL, Power BI, Excel, Machine Learning"
            

        # output path
        output_path = f"ResumeManager/OutputResumes/{company}_{job_role}_{today_date}/"
        # Create the directory if it doesn't exist
        os.makedirs(output_path, exist_ok=True)
        
        cover_input_dic = {
            "[job role]": job_role,
            "[company name]": company,
            "[company location]": company_location,
            "[job role]": job_role,
            "[date]": today_date,
            "[why company]": why_company,
            "[why me]": why_me
        }
        
        self.create_resume(resume_doc, job_role, top_skills, output_path)
        self.create_cover(cover_doc,output_path, cover_input_dic)
        
        
    def create_cover(self, cover_doc, output_path, cover_input_dic):
        self.table_edit_replace(cover_doc, "[job role]", cover_input_dic["[job role]"])
        for key, value in cover_input_dic.items():
            self.paragraph_edit_replace(cover_doc, key, value)
        docx_path = os.path.join(output_path, "Krishnakumar Cover Letter.docx")
        cover_doc.save(docx_path)
        self.save_to_pdf(output_path, docx_path)
        
    
    def create_resume(self, resume_doc, job_role, top_skills, output_path):
        self.table_edit_replace(resume_doc, "[job role]", job_role)
        self.paragraph_edit_replace(resume_doc, "[top skills]", top_skills)
        docx_path = os.path.join(output_path, "Krishnakumar Resume.docx")
        resume_doc.save(docx_path)
        self.save_to_pdf(output_path, docx_path)
        
        
    def save_to_pdf(self, output_path, docx_path):
        # Convert .docx to .pdf using LibreOffice
        libreoffice_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        # resume_pdf_path = os.path.join(output_path, f"Krishnakumar {doc_type}.pdf")
        try:
            subprocess.run([
                libreoffice_path,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                output_path,
                docx_path
            ], check=True)
        except subprocess.CalledProcessError as e:
            print(f"Error converting to PDF: {e}")